# coding: utf-8
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import Select
from xlwt import Workbook
import docx
from docx.enum.dml import MSO_THEME_COLOR
import time,chromedriver_binary,re

URL = 'https://portal.sa.dendai.ac.jp/up/faces/login/Com00505A.jsp'
ID = input('ID')
PASS = input('PASS')
seazon = input("前期:0 後期:1")
if seazon == "0":
    Seazon = "1年前期"
else: Seazon = "1年後期"
print(Seazon)
#ブラウザ起動
options = Options()
options.binary_location = "C:\Program Files\Google\Chrome\Application\chrome.exe"
#options.add_argument("--headless")
driver = webdriver.Chrome(options=options)
driver.get(URL)
wait = WebDriverWait(driver, 60)

#Excel初期化
wb = Workbook()
ws = wb.add_sheet(Seazon)
ws.write(0,0,Seazon) 
ws.write(0,1,"授業コード")
ws.write(0,2,"授業名")
ws.write(0,3,"英文名")
ws.write(0,4,"開講年度学期")
ws.write(0,5,"単位")
ws.write(0,6,"教室")
ws.write(0,7,"担当教員")
ws.write(0,8,"授業形式")
ws.write(0,9,"遠隔授業方法")
ws.write(0,10,"目的概要")
ws.write(0,11,"達成目標")
ws.write(0,12,"関連科目")
ws.write(0,13,"履修条件")
ws.write(0,14,"教科書名")
ws.write(0,15,"参考署名")
ws.write(0,16,"評価方法")
ws.write(0,17,"学習教育目標との対応")
ws.write(0,18,"DPとのの対応")
ws.write(0,19,"アクティブラーニングの実施")
ws.write(0,20,"ICTの活用")
ws.write(0,21,"実践的活用")
ws.write(0,22,"自由記載欄")
ws.write(0,23,"自由記載欄")
ws.write(0,24,"第1回")
ws.write(0,25,"第2回")
ws.write(0,26,"第3回")
ws.write(0,27,"第4回")
ws.write(0,28,"第5回")
ws.write(0,29,"第6回")
ws.write(0,30,"第7回")
ws.write(0,31,"第8回")
ws.write(0,32,"第9回")
ws.write(0,33,"第10回")
ws.write(0,34,"第11回")
ws.write(0,35,"第12回")
ws.write(0,36,"第13回")
ws.write(0,37,"第14回")
ws.write(0,38,"Email")
ws.write(0,39,"質問への対応")
ws.write(0,40,"履修上の注意事項 クラス分け")
ws.write(0,41,"履修上の注意事項　ガイダンス")
ws.write(0,42,"学習上の助言")

# ワードドキュメント作成
doc = docx.Document()

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR.HYPERLINK
    r.font.underline = True

#ログイン
driver.find_element_by_id('loginForm:userId').send_keys(ID)
driver.find_element_by_id('loginForm:password').send_keys(PASS)
driver.find_element_by_id('loginForm:loginButton').click()

print('ログイン成功')
#actions = ActionChains(driver)
#actions.move_to_element(driver.find_element_by_class_name("ui-menuitem-text")).perform()
# men
#u = driver.find_element_by_id('menu5')
# driver.getMouse().mouseMove(menu)
wait.until(EC.visibility_of_element_located((By.XPATH,'//*[@id="menuForm:mainMenu"]/ul/li[3]'))).click()
driver.find_element_by_xpath('//*[@id="menuForm:mainMenu"]/ul/li[3]/ul/table/tbody/tr/td[1]/ul/li[3]').click()

print('時間割ページ')
yobi = ["月","火","水","木","金","土"]
clasnum=0
res = ""
for i in range(1,48):
    try:
        clasnum+=1
        jigen = divmod(i,8)
        JIGEN = "{}曜日{}時間目".format(yobi[jigen[0]],jigen[1])
        if jigen[1] == 0:
            JIGEN = "{}曜日8時間目".format(yobi[jigen[0]-1])
        print(str(i)+":"+JIGEN)
        time.sleep(2)
        Jpath = '//*[@id="funcForm:j_idt250:'+str(seazon)+':j_idt259:'+str(jigen[1]-1)+':j_idt263:'+str(jigen[0])+':j_idt267:0:j_idt287"]' 
        if jigen[1] == 0:
            Jpath = '//*[@id="funcForm:j_idt250:'+str(seazon)+':j_idt259:7:j_idt263:'+str(jigen[0])+':j_idt267:0:j_idt287"]' 
        driver.find_element_by_xpath(Jpath).click()
        print("nowCaptureing"+str(i))
        #//*[@id="pkx02301:ch:table"]/div[3]/div[2]/div 授業コード
        #//*[@id="pkx02301:ch:table"]/div[4]/div[2]
        #//*[@id="pkx02301:ch:table"]/div[3]/div[4]
        #//*[@id="pkx02301:ch:table"]/div[29]/div[2]/div/div/div目的概要
        #//*[@id="pkx02301:ch:table"]/div[11]/div[2]/div
        #//*[@id="pkx02301:ch:table"]/div[43]/div[2]/div
        #//*[@id="pkx02301:ch:table"]/div[47]/div[2]/div
        timewait = wait.until(EC.visibility_of_element_located((By.XPATH,'//*[@id="pkx02301:ch:table"]/div[3]/div[2]/div')))
        print(timewait.text)
        number = 0
        ws.write(clasnum,number,JIGEN)
        doc.add_heading(JIGEN ,1)
        for i in range(3,50,1):

            number+=1
            try:
                atd = driver.find_element_by_xpath('//*[@id="pkx02301:ch:table"]/div[' + str(i) +']/div[2]/div')
                print(atd.text)
                ws.write(clasnum,number,atd.text)
            except:
                if number == 10 or 11:
                    number -= 1
                    pass
                else:
                    ws.write(clasnum,number,"")
            if number == 2:
                UrlText = atd.text
                doc.add_paragraph(UrlText)
            if i == 12:
                pattern = "https?://[\w/:%#\$&\?\(\)~\.=\+\-]+"
                url = re.findall(pattern, atd.text)
                for i in range(len(url)):    
                    p = doc.add_paragraph()
                    add_hyperlink(p,url[i],url[i])
        driver.find_element_by_xpath('//*[@id="pkx02301:dialog"]/div[1]/a[1]/span').click()
    except Exception as e:
        print("None")
        print(e)
        clasnum -= 1
        pass
wb.save(Seazon +".xls")
doc.save(Seazon +".docx")
driver.quit()